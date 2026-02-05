"""
OCR processing using Google Cloud Document AI and Vision API.
Handles images, PDFs, Excel, Word documents - extracts text for further processing.
Document AI is used for PDFs and structured documents, Vision API for simple images.
"""

import os
import logging
from typing import Optional, Tuple
from google.cloud import vision
from google.cloud import documentai
from dotenv import load_dotenv
import pandas as pd
import openpyxl
from docx import Document

load_dotenv()

logger = logging.getLogger(__name__)

# Document AI configuration
DOCUMENT_AI_PROJECT_ID = os.getenv('DOCUMENT_AI_PROJECT_ID')
DOCUMENT_AI_LOCATION = os.getenv('DOCUMENT_AI_LOCATION', 'us')  # e.g., 'us' or 'eu'
DOCUMENT_AI_PROCESSOR_ID = os.getenv('DOCUMENT_AI_PROCESSOR_ID')  # Optional: specific processor
USE_DOCUMENT_AI_FOR_IMAGES = os.getenv('USE_DOCUMENT_AI_FOR_IMAGES', 'true').lower() == 'true'


def detect_file_type(filename: str) -> str:
    """Detect file type from filename extension"""
    ext = filename.lower().split('.')[-1]
    
    image_types = ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff', 'webp']
    if ext in image_types:
        return 'image'
    elif ext == 'pdf':
        return 'pdf'
    elif ext in ['xlsx', 'xls']:
        return 'excel'
    elif ext == 'docx':
        return 'word'
    elif ext == 'csv':
        return 'csv'
    else:
        return 'unknown'


def extract_text_with_document_ai(file_path: str, mime_type: str) -> str:
    """
    Extract text using Google Cloud Document AI.
    Works with PDFs, images, and other document types.
    
    Args:
        file_path: Path to the file
        mime_type: MIME type (e.g., 'application/pdf', 'image/jpeg')
    
    Returns:
        Extracted text as string
    """
    try:
        if not DOCUMENT_AI_PROJECT_ID:
            raise ValueError("DOCUMENT_AI_PROJECT_ID not set in environment variables")
        
        # Initialize Document AI client
        client = documentai.DocumentProcessorServiceClient()
        
        # Construct processor name
        # If processor ID is provided, use it; otherwise use default
        if DOCUMENT_AI_PROCESSOR_ID:
            processor_name = client.processor_path(
                DOCUMENT_AI_PROJECT_ID,
                DOCUMENT_AI_LOCATION,
                DOCUMENT_AI_PROCESSOR_ID
            )
        else:
            # Use the default OCR processor
            processor_name = client.common_location_path(
                DOCUMENT_AI_PROJECT_ID,
                DOCUMENT_AI_LOCATION
            )
            # For default, we'll use the general document processor
            # You may need to create a processor in Google Cloud Console
            raise ValueError(
                "DOCUMENT_AI_PROCESSOR_ID is required. "
                "Create a processor in Google Cloud Console and set DOCUMENT_AI_PROCESSOR_ID in .env"
            )
        
        # Read the file
        with open(file_path, 'rb') as file:
            file_content = file.read()
        
        # Create the raw document
        raw_document = documentai.RawDocument(
            content=file_content,
            mime_type=mime_type
        )
        
        # Configure the process request
        request = documentai.ProcessRequest(
            name=processor_name,
            raw_document=raw_document
        )
        
        # Process the document
        result = client.process_document(request=request)
        document = result.document
        
        # Extract text from the document
        text = document.text
        
        logger.info(f"Document AI extracted {len(text)} characters")
        return text
        
    except Exception as e:
        logger.error(f"Error extracting text with Document AI: {e}")
        raise


def extract_text_from_image(image_path: str) -> str:
    """
    Extract text from image.
    Uses Document AI if enabled, otherwise falls back to Vision API.
    """
    # Try Document AI first if enabled
    if USE_DOCUMENT_AI_FOR_IMAGES and DOCUMENT_AI_PROJECT_ID:
        try:
            # Determine MIME type from file extension
            ext = image_path.lower().split('.')[-1]
            mime_map = {
                'jpg': 'image/jpeg',
                'jpeg': 'image/jpeg',
                'png': 'image/png',
                'gif': 'image/gif',
                'bmp': 'image/bmp',
                'tiff': 'image/tiff',
                'webp': 'image/webp'
            }
            mime_type = mime_map.get(ext, 'image/jpeg')
            
            return extract_text_with_document_ai(image_path, mime_type)
        except Exception as e:
            logger.warning(f"Document AI failed for image, falling back to Vision API: {e}")
            # Fall through to Vision API
    
    # Fallback to Vision API
    try:
        client = vision.ImageAnnotatorClient()
        
        with open(image_path, 'rb') as image_file:
            content = image_file.read()
        
        image = vision.Image(content=content)
        response = client.text_detection(image=image)
        
        if response.error.message:
            raise Exception(f"Google Vision API error: {response.error.message}")
        
        texts = response.text_annotations
        if texts:
            # First annotation contains all detected text
            return texts[0].description
        else:
            return ""
            
    except Exception as e:
        logger.error(f"Error extracting text from image: {e}")
        raise


def extract_text_from_pdf(pdf_path: str) -> str:
    """
    Extract text from PDF using Google Cloud Document AI.
    Document AI provides much better PDF handling than Vision API.
    """
    try:
        return extract_text_with_document_ai(pdf_path, 'application/pdf')
    except Exception as e:
        logger.error(f"Error extracting text from PDF with Document AI: {e}")
        # Fallback to Vision API (limited functionality)
        logger.warning("Falling back to Vision API for PDF (limited functionality)")
        try:
            client = vision.ImageAnnotatorClient()
            
            with open(pdf_path, 'rb') as pdf_file:
                content = pdf_file.read()
            
            # Note: Vision API doesn't handle PDFs well - this is a fallback
            # For multi-page PDFs, you'd need to convert pages to images first
            image = vision.Image(content=content)
            response = client.text_detection(image=image)
            
            if response.error.message:
                raise Exception(f"Google Vision API error: {response.error.message}")
            
            texts = response.text_annotations
            if texts:
                return texts[0].description
            else:
                return ""
        except Exception as fallback_error:
            logger.error(f"Fallback Vision API also failed: {fallback_error}")
            raise


def extract_text_from_excel(excel_path: str) -> str:
    """Extract text from Excel file"""
    try:
        # Read Excel file
        df = pd.read_excel(excel_path)
        
        # Convert DataFrame to text representation
        # Combine all cells into a single text string
        text_lines = []
        for _, row in df.iterrows():
            row_text = ' '.join([str(val) for val in row.values if pd.notna(val)])
            if row_text.strip():
                text_lines.append(row_text)
        
        return '\n'.join(text_lines)
        
    except Exception as e:
        logger.error(f"Error extracting text from Excel: {e}")
        raise


def extract_text_from_word(word_path: str) -> str:
    """Extract text from Word document"""
    try:
        doc = Document(word_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        return '\n'.join(paragraphs)
        
    except Exception as e:
        logger.error(f"Error extracting text from Word document: {e}")
        raise


def extract_text_from_csv(csv_path: str) -> str:
    """Extract text from CSV file"""
    try:
        df = pd.read_csv(csv_path)
        
        # Convert DataFrame to text representation
        text_lines = []
        for _, row in df.iterrows():
            row_text = ' '.join([str(val) for val in row.values if pd.notna(val)])
            if row_text.strip():
                text_lines.append(row_text)
        
        return '\n'.join(text_lines)
        
    except Exception as e:
        logger.error(f"Error extracting text from CSV: {e}")
        raise


def process_file(file_path: str, filename: str) -> Tuple[str, str]:
    """
    Process uploaded file and extract text.
    Returns tuple: (extracted_text, file_type)
    """
    file_type = detect_file_type(filename)
    logger.info(f"Processing file: {filename}, type: {file_type}")
    
    text = ""
    
    if file_type == 'image':
        text = extract_text_from_image(file_path)
    elif file_type == 'pdf':
        text = extract_text_from_pdf(file_path)
    elif file_type == 'excel':
        text = extract_text_from_excel(file_path)
    elif file_type == 'word':
        text = extract_text_from_word(file_path)
    elif file_type == 'csv':
        text = extract_text_from_csv(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
    
    logger.info(f"Extracted {len(text)} characters from {filename}")
    return text, file_type
